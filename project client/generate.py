import sqlite3
from docx import Document
import os
from petrovich.main import Petrovich
from petrovich.enums import Case, Gender


def get_gender(middle_name: str) -> Gender:
    """Определяем пол по отчеству (упрощённо)."""
    if middle_name.endswith("ич"):
        return Gender.MALE
    elif middle_name.endswith("на"):
        return Gender.FEMALE
    return Gender.MALE


def personalize_name(last_name: str, first_name: str, middle_name: str):
    """Возвращает два варианта ФИО: полное и имя+отчество в дательном падеже."""
    p = Petrovich()
    gender = get_gender(middle_name)

    # Полное ФИО
    name_full = f"{last_name} {first_name} {middle_name}"
    # Имя + Отчество в дательном падеже
    first_d = p.firstname(first_name, Case.DATIVE, gender)
    middle_d = p.middlename(middle_name, Case.DATIVE, gender)
    name_short = f"{first_d} {middle_d}"

    return name_full, name_short


def replace_placeholders(doc: Document, name_full: str, name_short: str):
    """Заменяем плейсхолдеры {name_full} и {name_short} во всех параграфах и таблицах."""
    for para in doc.paragraphs:
        para.text = para.text.replace("{name_full}", name_full)
        para.text = para.text.replace("{name_short}", name_short)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("{name_full}", name_full)
                cell.text = cell.text.replace("{name_short}", name_short)


def generate_letters(db_path: str, template_path: str, output_dir="letters"):
    """Создание персональных писем на основе шаблона и базы клиентов."""
    if not os.path.exists(db_path):
        print(f"❌ Файл БД не найден: {db_path}")
        return

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Проверяем наличие таблицы vip_clients
    cursor.execute("""
        SELECT name FROM sqlite_master 
        WHERE type='table' AND name='vip_clients'
    """)
    if cursor.fetchone() is None:
        print("❌ Таблица 'vip_clients' не найдена в базе данных!")
        conn.close()
        return

    cursor.execute("SELECT last_name, first_name, middle_name FROM vip_clients")
    clients = cursor.fetchall()
    conn.close()

    if not clients:
        print("❌ В таблице 'vip_clients' нет клиентов.")
        return

    os.makedirs(output_dir, exist_ok=True)

    for i, (last_name, first_name, middle_name) in enumerate(clients, start=1):
        doc = Document(template_path)
        name_full, name_short = personalize_name(last_name, first_name, middle_name)

        replace_placeholders(doc, name_full, name_short)

        out_path = os.path.join(output_dir, f"letter_{i}_{last_name}.docx")
        doc.save(out_path)
        print(f"✅ Создано письмо: {out_path}")


if __name__ == "__main__":
    # Путь к файлу БД
    db_path = r"C:\Users\yanin\TUSUR\tusur\project client\clients.db"
    # Путь к шаблону Word
    template_path = r"C:\Users\yanin\TUSUR\tusur\project client\template.docx"

    if not os.path.exists(template_path):
        print(f"❌ Шаблон не найден: {template_path}")
    else:
        generate_letters(db_path, template_path)
